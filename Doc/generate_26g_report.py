"""Generate the editable G26 report and its figures from the Markdown draft."""

from pathlib import Path
import math
import re

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "26G_设计报告_第一版.md"
OUTPUT = ROOT / "26G_设计报告_第一版.docx"
ASSETS = ROOT / "26G_设计报告_第一版_assets"
FONT_REGULAR = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\msyhbd.ttc")
COE = ROOT.parent / "25G_PL" / "script" / "fir_26g_decim3_q17.coe"

INK = "203047"
BLUE = "2D6A8A"
TEAL = "2A7F78"
GOLD = "B98224"
RED = "B54747"
PALE_BLUE = "EAF2F7"
PALE_TEAL = "E8F3F1"
PALE_GOLD = "F7F0E2"
GRID = "CFD8DE"


def pil_font(size, bold=False):
    return ImageFont.truetype(str(FONT_BOLD if bold else FONT_REGULAR), size)


def text_center(draw, box, text, size=30, bold=False, fill="#203047", spacing=6):
    font = pil_font(size, bold)
    x0, y0, x1, y1 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=spacing)
    x = (x0 + x1 - (bbox[2] - bbox[0])) / 2
    y = (y0 + y1 - (bbox[3] - bbox[1])) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, align="center", spacing=spacing)


def box(draw, xy, text, fill="#FFFFFF", outline="#2D6A8A", size=28):
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    text_center(draw, xy, text, size=size, bold=True)


def arrow(draw, start, end, fill="#62727C", width=5):
    draw.line((start, end), fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    for delta in (2.55, -2.55):
        p = (end[0] + length * math.cos(angle + delta), end[1] + length * math.sin(angle + delta))
        draw.line((end, p), fill=fill, width=width)


def save_canvas(image, name):
    image.save(ASSETS / name, dpi=(180, 180))


def make_system_architecture():
    image = Image.new("RGB", (1900, 790), "white")
    draw = ImageDraw.Draw(image)
    draw.text((55, 30), "周期信号测量分析装置端到端架构", font=pil_font(42, True), fill="#203047")
    draw.rounded_rectangle((40, 110, 810, 710), 22, fill="#F7F0E2", outline="#D8C79F", width=2)
    draw.rounded_rectangle((825, 110, 1395, 710), 22, fill="#EAF2F7", outline="#AFC8D7", width=2)
    draw.rounded_rectangle((1410, 110, 1860, 710), 22, fill="#E8F3F1", outline="#A8CDC8", width=2)
    draw.text((65, 130), "模拟前端", font=pil_font(30, True), fill="#8A641F")
    draw.text((850, 130), "PL 数据面", font=pil_font(30, True), fill="#2D6A8A")
    draw.text((1435, 130), "PS 算法与显示接口", font=pil_font(30, True), fill="#2A7F78")

    items = [
        ((70, 260, 255, 430), "信号源\n50 Ω 模式", "#FFFDF7", "#B98224"),
        ((280, 260, 465, 430), "板端 50 Ω\n对地终端", "#FFFDF7", "#B98224"),
        ((490, 260, 675, 430), "约五阶 LPF\nfc≈800 kHz", "#FFFDF7", "#B98224"),
        ((700, 260, 785, 430), "AD9226\n5.12006\nMSPS", "#FFFDF7", "#B98224"),
        ((855, 230, 1080, 460), "码型转换\n39-tap FIR /3\n舍入与饱和", "#F8FBFD", "#2D6A8A"),
        ((1110, 260, 1265, 430), "AXIS FIFO\n4096 点\nTLAST", "#F8FBFD", "#2D6A8A"),
        ((1290, 260, 1370, 430), "SG\nDMA", "#F8FBFD", "#2D6A8A"),
        ((1440, 220, 1830, 470), "4096 点 RFFT 候选\n整数谐波联合最小二乘\n频率细化 + BIC\n幅值校准 + 指标重建", "#F8FCFB", "#2A7F78"),
    ]
    for xy, text, fill, outline in items:
        box(draw, xy, text, fill, outline, 24 if "AD9226" not in text else 21)
    for first, second in zip(items, items[1:]):
        arrow(draw, ((first[0][2], (first[0][1] + first[0][3]) / 2)),
              ((second[0][0], (second[0][1] + second[0][3]) / 2)))
    draw.text((860, 525), "Fs = 1.7066867 MSPS    ·    signed16    ·    8192 byte/frame",
              font=pil_font(25), fill="#2D6A8A")
    draw.text((1450, 525), "Upp / Urms / 基频 / 正频率分量",
              font=pil_font(25), fill="#2A7F78")
    draw.text((1450, 575), "1、3 周期各 640 点稳定波形",
              font=pil_font(25), fill="#2A7F78")
    save_canvas(image, "system_architecture.png")


def read_coefficients():
    text = COE.read_text(encoding="utf-8")
    body = text.split("coefdata=", 1)[1]
    return np.array([int(x) for x in re.findall(r"-?\d+", body)], dtype=float) / (2**17)


def make_fir_response():
    coefficients = read_coefficients()
    fs = 5_120_060.0
    freq = np.linspace(0, fs / 2, 5000)
    n = np.arange(len(coefficients))
    response = np.exp(-2j * np.pi * np.outer(freq / fs, n)) @ coefficients
    db = np.maximum(20 * np.log10(np.maximum(np.abs(response), 1e-6)), -100)

    image = Image.new("RGB", (1800, 980), "white")
    draw = ImageDraw.Draw(image)
    plot = (165, 100, 1710, 820)
    draw.text((60, 25), "39-tap Q1.17 抽取 FIR 量化频率响应", font=pil_font(40, True), fill="#203047")
    for y_db in range(-100, 1, 20):
        y = plot[3] - (y_db + 100) / 100 * (plot[3] - plot[1])
        draw.line((plot[0], y, plot[2], y), fill="#D9E0E4", width=2)
        draw.text((75, y - 15), str(y_db), font=pil_font(24), fill="#52636D")
    for x_mhz in (0, 0.5, 1.0, 1.5, 2.0, 2.56):
        x = plot[0] + x_mhz / (fs / 2 / 1e6) * (plot[2] - plot[0])
        draw.line((x, plot[1], x, plot[3]), fill="#E7ECEF", width=2)
        draw.text((x - 28, plot[3] + 18), f"{x_mhz:g}", font=pil_font(24), fill="#52636D")
    points = []
    for f, value in zip(freq[::3], db[::3]):
        x = plot[0] + f / (fs / 2) * (plot[2] - plot[0])
        y = plot[3] - (value + 100) / 100 * (plot[3] - plot[1])
        points.append((x, y))
    draw.line(points, fill="#2D6A8A", width=5)
    for f_hz, label, color in ((500_000, "500 kHz\n-0.008 dB", "#2A7F78"),
                               (1_000_000, "1 MHz\n-67.63 dB", "#B54747")):
        idx = int(f_hz / (fs / 2) * (len(freq) - 1))
        x = plot[0] + f_hz / (fs / 2) * (plot[2] - plot[0])
        y = plot[3] - (db[idx] + 100) / 100 * (plot[3] - plot[1])
        draw.ellipse((x - 8, y - 8, x + 8, y + 8), fill=color)
        draw.multiline_text((x + 15, y - 60), label, font=pil_font(24, True), fill=color, spacing=3)
    draw.line((plot[0], plot[1], plot[0], plot[3]), fill="#52636D", width=3)
    draw.line((plot[0], plot[3], plot[2], plot[3]), fill="#52636D", width=3)
    draw.text((25, 410), "幅度/dB", font=pil_font(27, True), fill="#203047")
    draw.text((830, 885), "频率/MHz（原始采样域）", font=pil_font(27, True), fill="#203047")
    save_canvas(image, "fir_response.png")


def make_analysis_flow():
    image = Image.new("RGB", (1800, 930), "white")
    draw = ImageDraw.Draw(image)
    draw.text((55, 30), "周期信号参数估计算法", font=pil_font(42, True), fill="#203047")
    steps = [
        "4096 点 signed16\n估计并去除中点偏置",
        "Hann 窗 + RFFT\n局部峰与频率插值",
        "构造整数谐波\n基频与阶次假设",
        "未加窗时域\n联合最小二乘",
        "4 轮 × 5 点\n残差频率细化",
        "BIC 选择\n2/3 分量模型",
        "3σ 与 0.5 mVpeak\n基波有效性门限",
        "频率补偿并重算\nUrms、Upp 与波形",
    ]
    positions = [(70 + i * 425, 160, 400 + i * 425, 360) for i in range(4)]
    positions += [(1345 - i * 425, 540, 1675 - i * 425, 740) for i in range(4)]
    colors = ["#EAF2F7"] * 4 + ["#E8F3F1"] * 4
    outlines = ["#2D6A8A"] * 4 + ["#2A7F78"] * 4
    for pos, step, fill, outline in zip(positions, steps, colors, outlines):
        box(draw, pos, step, fill, outline, 26)
    for i in range(3):
        arrow(draw, (positions[i][2], 260), (positions[i + 1][0], 260))
    arrow(draw, ((positions[3][0] + positions[3][2]) / 2, positions[3][3]),
          ((positions[4][0] + positions[4][2]) / 2, positions[4][1]))
    for i in range(4, 7):
        arrow(draw, (positions[i][0], 640), (positions[i + 1][2], 640))
    draw.rounded_rectangle((510, 800, 1290, 875), 16, fill="#F7F0E2", outline="#B98224", width=2)
    text_center(draw, (510, 800, 1290, 875), "FFT 负责定位，联合拟合负责最终定量", 27, True, "#805B1F")
    save_canvas(image, "analysis_flow.png")


def make_measurement_state():
    image = Image.new("RGB", (1900, 620), "white")
    draw = ImageDraw.Draw(image)
    draw.text((55, 30), "KEY1 触发的一次性测量状态机", font=pil_font(42, True), fill="#203047")
    labels = ["ARMED", "ALIGN\n丢弃至 TLAST", "WARMUP\n丢弃 3 帧", "CAPTURE\n8192 byte", "ANALYZE\n拟合与校准", "HOLD\n稳定结果"]
    positions = []
    width, gap = 250, 50
    for i, label in enumerate(labels):
        pos = (55 + i * (width + gap), 210, 55 + i * (width + gap) + width, 380)
        positions.append(pos)
        box(draw, pos, label, "#EAF2F7" if i < 4 else "#E8F3F1", "#2D6A8A" if i < 4 else "#2A7F78", 25)
    for a, b in zip(positions, positions[1:]):
        arrow(draw, (a[2], 295), (b[0], 295))
    draw.text((125, 410), "KEY1", font=pil_font(24, True), fill="#B98224")
    draw.arc((positions[-1][0], 420, positions[-1][2], 555), 0, 180, fill="#B54747", width=4)
    draw.line((positions[-1][0], 488, positions[0][2], 488), fill="#B54747", width=4)
    arrow(draw, (positions[0][2], 488), (positions[0][2], positions[0][3]), fill="#B54747", width=4)
    draw.text((820, 500), "KEY2：清除结果并重新装载", font=pil_font(25, True), fill="#B54747")
    save_canvas(image, "measurement_state.png")


def make_interference_delta():
    conditions = ["无干扰", "1.0 MHz", "1.5 MHz", "2.0 MHz"]
    series = {
        "H1": [0.000, 0.041, -0.054, 0.004],
        "H2": [0.000, -0.056, -0.022, 0.020],
        "Urms": [0.000, -0.007, -0.038, 0.012],
        "Upp": [0.000, -0.054, -0.116, 0.050],
    }
    colors = {"H1": "#2D6A8A", "H2": "#2A7F78", "Urms": "#B98224", "Upp": "#B54747"}
    image = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((55, 30), "高频干扰下有效参数相对基准的变化", font=pil_font(40, True), fill="#203047")
    plot = (170, 120, 1710, 730)
    y_min, y_max = -0.20, 0.20
    for value in np.arange(y_min, y_max + 0.001, 0.05):
        y = plot[3] - (value - y_min) / (y_max - y_min) * (plot[3] - plot[1])
        draw.line((plot[0], y, plot[2], y), fill="#D9E0E4", width=2)
        draw.text((65, y - 14), f"{value:+.2f}", font=pil_font(23), fill="#52636D")
    xs = np.linspace(plot[0] + 120, plot[2] - 120, len(conditions))
    for label, values in series.items():
        points = []
        for x, value in zip(xs, values):
            y = plot[3] - (value - y_min) / (y_max - y_min) * (plot[3] - plot[1])
            points.append((x, y))
            draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill=colors[label])
        draw.line(points, fill=colors[label], width=5)
    for x, condition in zip(xs, conditions):
        draw.text((x - 55, plot[3] + 25), condition, font=pil_font(23), fill="#52636D")
    draw.line((plot[0], plot[1], plot[0], plot[3]), fill="#52636D", width=3)
    draw.line((plot[0], plot[3], plot[2], plot[3]), fill="#52636D", width=3)
    draw.text((35, 385), "变化/mV", font=pil_font(26, True), fill="#203047")
    legend_x = 490
    for label in series:
        draw.line((legend_x, 820, legend_x + 50, 820), fill=colors[label], width=6)
        draw.text((legend_x + 62, 802), label, font=pil_font(24), fill="#203047")
        legend_x += 220
    save_canvas(image, "interference_delta.png")


def set_run_font(run, name="宋体", size=10.5, bold=None, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))
    set_run_font(run, size=9)


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    node = sect_pr.find(qn("w:pgNumType"))
    if node is None:
        node = OxmlElement("w:pgNumType")
        sect_pr.append(node)
    node.set(qn("w:start"), str(start))


def configure_document(document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0.74)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)

    for name, level, size in (("Heading 1", 1, 15), ("Heading 2", 2, 13), ("Heading 3", 3, 11.5)):
        style = document.styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(8 if level == 1 else 5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    for style_name, size, bold in (("Report Caption", 9, False), ("Report Formula", 10.5, False)):
        style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = style_name == "Report Caption"


def add_inline_runs(paragraph, text, size=10.5):
    text = text.replace("$", "")
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        elif part.startswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="等线", size=max(size - 0.5, 8.5), color="37474F")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def readable_formula(text):
    text = text.strip().strip("$")
    text = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", text)
    text = re.sub(r"\\sqrt\{([^{}]+)\}", r"√(\1)", text)
    text = re.sub(r"\\text\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\operatorname\{([^{}]+)\}", r"\1", text)
    replacements = (
        (r"\operatorname", ""), (r"\qquad", "    "), (r"\left", ""), (r"\right", ""),
        (r"\approx", "≈"), (r"\times", "×"), (r"\varphi", "φ"), (r"\Delta", "Δ"),
        (r"\sum", "Σ"), (r"\pm", "±"), (r"\sim", "～"), (r"\le", "≤"), (r"\ge", "≥"),
        (r"\pi", "π"), (r"\ln", "ln"), (r"\,", ""),
    )
    for old, new in replacements:
        text = text.replace(old, new)
    text = re.sub(r"\\([A-Za-z]+)", r"\1", text)
    text = re.sub(r"_\{([^{}]+)\}", r"_\1", text)
    text = re.sub(r"\^\{([^{}]+)\}", r"^\1", text)
    return text.replace("{", "").replace("}", "").replace("\\ ", " ")


def add_table(document, rows):
    data = []
    for row in rows:
        cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", cell.replace(" ", "")) for cell in cells):
            continue
        data.append(cells)
    if not data:
        return
    table = document.add_table(rows=len(data), cols=max(len(row) for row in data))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "DCE8EF")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_runs(paragraph, value, 8.5)
            if r_idx == 0:
                for run in paragraph.runs:
                    run.bold = True
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(document, text, level):
    paragraph = document.add_heading(level=level)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    add_inline_runs(paragraph, text, {1: 15, 2: 13, 3: 11.5}[level])


def render_cover(document, lines):
    document.add_paragraph().paragraph_format.space_after = Pt(65)
    for line in lines:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph().paragraph_format.space_after = Pt(10)
            continue
        if stripped.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped[2:])
            set_run_font(run, name="黑体", size=18, bold=True)
            paragraph.paragraph_format.space_after = Pt(45)
        elif stripped.startswith("## "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped[3:])
            set_run_font(run, name="黑体", size=24, bold=True)
            paragraph.paragraph_format.space_after = Pt(25)
        elif stripped.startswith("### "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped[4:])
            set_run_font(run, name="黑体", size=15, bold=True)
            paragraph.paragraph_format.space_after = Pt(80)
        else:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(14)
            run = paragraph.add_run(stripped)
            set_run_font(run, size=13, bold="2026 年" in stripped)
def render_markdown(document, lines):
    i = 0
    body_started = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "<!-- PAGEBREAK -->":
            document.add_page_break()
            i += 1
            continue
        if stripped == "<!-- BODY_START -->":
            body_started = True
            section = document.add_section(WD_SECTION.NEW_PAGE)
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.4)
            section.bottom_margin = Cm(2.2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.2)
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            header = section.header.paragraphs[0]
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(header, "2026 年全国大学生电子设计竞赛 · 周期信号测量分析装置", 8.5)
            add_page_number(section.footer.paragraphs[0])
            set_page_number_start(section, 1)
            i += 1
            continue
        image_match = re.fullmatch(r"!\[(.+)]\((.+)\)", stripped)
        if image_match:
            caption, path = image_match.groups()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            run = paragraph.add_run()
            run.add_picture(str(ROOT / path), width=Cm(16.0))
            cap = document.add_paragraph(caption, style="Report Caption")
            cap.paragraph_format.keep_with_next = False
            i += 1
            continue
        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            add_table(document, rows)
            continue
        if stripped.startswith("$$"):
            formula = stripped
            while not (formula.endswith("$$") and len(formula) > 4):
                i += 1
                formula += " " + lines[i].strip()
            paragraph = document.add_paragraph(style="Report Formula")
            add_inline_runs(paragraph, readable_formula(formula), 10.5)
            i += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            add_heading(document, heading.group(2), len(heading.group(1)))
            i += 1
            continue
        if re.match(r"^表\s*\d+", stripped):
            paragraph = document.add_paragraph(style="Report Caption")
            add_inline_runs(paragraph, stripped, 9)
            i += 1
            continue
        list_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if list_match:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.left_indent = Cm(0.74)
            add_inline_runs(paragraph, list_match.group(2))
            i += 1
            continue
        paragraph = document.add_paragraph()
        if body_started:
            add_inline_runs(paragraph, stripped)
        else:
            paragraph.paragraph_format.line_spacing = 1.08
            paragraph.paragraph_format.first_line_indent = Cm(0.63)
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_runs(paragraph, stripped, 9)
        i += 1


def generate_document():
    ASSETS.mkdir(exist_ok=True)
    make_system_architecture()
    make_fir_response()
    make_analysis_flow()
    make_measurement_state()
    make_interference_delta()

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    marker = lines.index("<!-- COVER_END -->")
    document = Document()
    configure_document(document)
    render_cover(document, lines[:marker])
    render_markdown(document, lines[marker + 1 :])
    document.core_properties.title = "周期信号测量分析装置（G题）设计报告第一版"
    document.core_properties.subject = "2026年全国大学生电子设计竞赛"
    document.core_properties.author = "参赛队（待填写）"
    document.core_properties.keywords = "Zynq, 周期信号, 谐波估计, FIR, 最小二乘"
    document.save(OUTPUT)

    check = Document(OUTPUT)
    assert len(check.paragraphs) > 100
    assert len(check.tables) >= 6
    assert all((ASSETS / name).exists() for name in (
        "system_architecture.png", "fir_response.png", "analysis_flow.png",
        "measurement_state.png", "interference_delta.png"
    ))
    print(f"generated: {OUTPUT}")
    print(f"paragraphs={len(check.paragraphs)} tables={len(check.tables)} figures=5")


if __name__ == "__main__":
    generate_document()
